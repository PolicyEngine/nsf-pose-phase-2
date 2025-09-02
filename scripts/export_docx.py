#!/usr/bin/env python3
"""
Export the assembled application markdown to a .docx for page review.
This is a best-effort conversion that maps headings and paragraphs and
includes basic list support. It does not guarantee pagination identical
to PDF/Research.gov, but gives a working document to check approximate
length.
"""
from __future__ import annotations

from pathlib import Path
from docx import Document  # type: ignore
from docx.shared import Inches, Pt  # type: ignore

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / 'docs' / 'pose' / 'assembled.md'
DOCX_OUT = ROOT / 'docs' / 'pose' / 'assembled.docx'


def main() -> None:
    if not MD_PATH.exists():
        print(f"Missing assembled markdown: {MD_PATH}")
        return
    lines = MD_PATH.read_text(encoding='utf-8').splitlines()
    doc = Document()
    # Set margins to ~1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    # Base font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    in_code = False
    in_list = False
    for raw in lines:
        line = raw.rstrip('\n')
        if line.startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            doc.add_paragraph(line)
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            in_list = False
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            in_list = False
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            in_list = False
        elif line.startswith('- '):
            para = doc.add_paragraph(line[2:].strip())
            para.style = doc.styles['List Bullet']
            in_list = True
        elif line.strip() == '---':
            doc.add_page_break()
            in_list = False
        elif line.strip() == '':
            # Blank line: paragraph break
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)
            in_list = False

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(f"✅ Wrote {DOCX_OUT}")


if __name__ == '__main__':
    main()

